import os
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from langchain.schema import Document
from services.retrieval import Retriever

embedder = SentenceTransformer('all-MiniLM-L6-v2')

def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    texts = []
    for p in reader.pages:
        try:
            texts.append(p.extract_text() or '')
        except Exception:
            continue
    return '\n'.join(texts)

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    return '\n'.join([p.text for p in doc.paragraphs])

def ingest_file(path: str, retriever: Retriever):
    """Ingest a single file and add to retriever index."""
    _, ext = os.path.splitext(path.lower())
    if ext == '.pdf':
        text = extract_text_from_pdf(path)
    elif ext in ('.docx', '.doc'):
        text = extract_text_from_docx(path)
    else:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

    # naive chunking
    chunks = []
    chunk_size = 800
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i+chunk_size]
        chunks.append(Document(page_content=chunk, metadata={'source': os.path.basename(path), 'pos': i}))

    embeddings = [embedder.encode(d.page_content).tolist() for d in chunks]
    retriever.add_documents(chunks, embeddings)
